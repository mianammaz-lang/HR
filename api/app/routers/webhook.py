"""
Webhook endpoint — push CVs into the system via API key.

Supports:
  - Multipart file upload (PDF, DOCX, TXT)
  - JSON payload with base64-encoded CV content
  - JSON payload with raw text content

Auth: X-API-Key header (configured in Settings > Webhook).

Usage:
  POST /api/webhook/cv
  Header: X-API-Key: your-webhook-secret
  Body (multipart):
    file=<cv.pdf>&source=linkedin&department=Engineering
  Body (JSON):
    {"content": "<base64 or raw text>", "filename": "cv.pdf", "source": "linkedin"}
"""
import os
import uuid
import hashlib
from fastapi import APIRouter, Depends, HTTPException, Header, UploadFile, File, Form, Request
from fastapi.responses import JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from typing import Optional
from pydantic import BaseModel
import base64

from app.database import get_db
from app.models import (
    Candidate, Document, Skill, EmploymentHistory, SyncStatusRecord,
    AuditLog, CareerLevel, SystemSetting
)
from app.services.cv_parser import BuiltInCVParser
from app.config import UPLOAD_DIR

router = APIRouter(prefix="/api/webhook", tags=["Webhook"])
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ─── Helpers ──────────────────────────────────────────────────────────────────

async def _get_webhook_key(db: AsyncSession) -> Optional[str]:
    result = await db.execute(
        select(SystemSetting).where(SystemSetting.key == "webhook_api_key")
    )
    row = result.scalar_one_or_none()
    return row.value if row else None


async def _verify_api_key(x_api_key: Optional[str], db: AsyncSession):
    """Verify the webhook API key. If no key is configured, allow all (dev mode)."""
    stored_key = await _get_webhook_key(db)
    if not stored_key:
        return  # No key configured — open mode (dev/testing)
    if not x_api_key or x_api_key != stored_key:
        raise HTTPException(status_code=401, detail="Invalid or missing X-API-Key header")


async def _extract_text(content: bytes, filename: str) -> str:
    """Extract text from file bytes."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    if ext == "pdf":
        try:
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            return text
        except Exception:
            return content.decode("utf-8", errors="ignore")

    elif ext in ("docx", "doc"):
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            return content.decode("utf-8", errors="ignore")

    else:
        return content.decode("utf-8", errors="ignore")


async def _create_candidate_from_cv(
    cv_text: str,
    filename: str,
    source: Optional[str],
    department: Optional[str],
    db: AsyncSession,
) -> Candidate:
    """Parse CV text and create a full candidate record. Returns the candidate."""
    parser = BuiltInCVParser()
    parsed = parser.parse(cv_text)

    # Optional AI enhancement
    from app.config import settings
    if settings.OPENROUTER_API_KEY:
        try:
            from app.services.ai_service import AIService
            ai = AIService()
            ai_parsed = await ai.parse_cv(cv_text)
            for key in ("full_name", "email", "phone", "linkedin_url", "city"):
                if ai_parsed.get(key) and (not parsed.get(key) or parsed[key] == "Unknown"):
                    parsed[key] = ai_parsed[key]
            if ai_parsed.get("skills") and len(ai_parsed["skills"]) > len(parsed.get("skills", [])):
                parsed["skills"] = ai_parsed["skills"]
            if ai_parsed.get("employment_history"):
                parsed["employment_history"] = ai_parsed["employment_history"]
        except Exception:
            pass

    # Override department if specified
    if department:
        parsed["department_tag"] = department

    # Duplicate hash
    email = (parsed.get("email") or "").strip().lower()
    phone = (parsed.get("phone") or "").strip()
    raw_hash = (email + phone) if (email or phone) else ""
    dup_hash = hashlib.sha256(raw_hash.encode()).hexdigest() if raw_hash else None

    is_duplicate_of = None
    if dup_hash:
        existing = await db.execute(
            select(Candidate).where(Candidate.duplicate_check_hash == dup_hash)
        )
        ex = existing.scalar_one_or_none()
        if ex:
            is_duplicate_of = ex.candidate_id

    # Career level
    cl_map = {
        "entry": "Entry", "junior": "Entry", "mid": "Mid", "intermediate": "Mid",
        "senior": "Senior", "sr": "Senior", "lead": "Lead",
        "managerial": "Managerial", "manager": "Managerial", "director": "Managerial",
    }
    try:
        cl = CareerLevel(cl_map.get((parsed.get("career_level") or "Mid").lower(), "Mid"))
    except ValueError:
        cl = CareerLevel.mid

    # Source channel mapping
    source_map = {
        "linkedin": "LinkedIn", "indeed": "Indeed", "referral": "Referral",
        "agency": "Agency", "website": "Website", "direct": "Direct",
        "email": "Direct", "webhook": "Direct",
    }
    source_enum = None
    if source:
        from app.models import SourceChannel
        try:
            source_enum = SourceChannel(source_map.get(source.lower(), source))
        except ValueError:
            source_enum = None

    candidate = Candidate(
        full_name=parsed.get("full_name", "Unknown"),
        email=parsed.get("email"),
        phone=parsed.get("phone"),
        linkedin_url=parsed.get("linkedin_url"),
        city=parsed.get("city"),
        source_channel=source_enum,
        applied_designation=parsed.get("applied_designation"),
        department_tag=parsed.get("department_tag"),
        career_level=cl,
        tags=parsed.get("tags", []),
        duplicate_check_hash=dup_hash,
        is_duplicate_of=is_duplicate_of,
    )
    db.add(candidate)
    await db.flush()

    for skill_name in parsed.get("skills", []):
        db.add(Skill(candidate_id=candidate.candidate_id, skill_name=skill_name))

    for hist in parsed.get("employment_history", []):
        db.add(EmploymentHistory(
            candidate_id=candidate.candidate_id,
            company=hist.get("company", ""),
            title=hist.get("title", ""),
            duration=hist.get("duration", ""),
            description=hist.get("description", ""),
        ))

    db.add(Document(
        candidate_id=candidate.candidate_id,
        resume_file_url=f"webhook:{filename}",
        original_filename=filename,
    ))
    db.add(SyncStatusRecord(candidate_id=candidate.candidate_id))

    await db.flush()
    return candidate


# ─── JSON Request Schema ──────────────────────────────────────────────────────

class WebhookCVRequest(BaseModel):
    content: str  # base64-encoded or raw text
    filename: str = "cv.txt"
    source: Optional[str] = None
    department: Optional[str] = None
    format: str = "text"  # "text" or "base64"


# ─── Webhook Endpoints ────────────────────────────────────────────────────────

@router.post("/cv")
async def receive_cv_webhook(
    request: Request,
    file: Optional[UploadFile] = File(None),
    content: Optional[str] = Form(None),
    filename: Optional[str] = Form(None),
    source: Optional[str] = Form(None),
    department: Optional[str] = Form(None),
    x_api_key: Optional[str] = Header(None),
    db: AsyncSession = Depends(get_db),
):
    """
    Receive a CV via webhook.

    **Option 1 — Multipart file upload:**
    ```
    POST /api/webhook/cv
    Headers: X-API-Key: your-key
    Body: multipart/form-data
      file: <cv.pdf>
      source: linkedin        (optional)
      department: Engineering (optional)
    ```

    **Option 2 — JSON with raw text:**
    ```
    POST /api/webhook/cv
    Headers: X-API-Key: your-key, Content-Type: application/json
    Body:
    {
      "content": "Full CV text here...",
      "filename": "cv.pdf",
      "source": "linkedin",
      "department": "Engineering",
      "format": "text"
    }
    ```

    **Option 3 — JSON with base64:**
    ```
    POST /api/webhook/cv
    Headers: X-API-Key: your-key, Content-Type: application/json
    Body:
    {
      "content": "<base64-encoded-file>",
      "filename": "cv.pdf",
      "source": "email",
      "format": "base64"
    }
    ```
    """
    await _verify_api_key(x_api_key, db)

    # Determine input mode
    cv_bytes = None
    cv_filename = "cv.txt"
    src = source
    dept = department

    content_type = request.headers.get("content-type", "")

    if "application/json" in content_type:
        # JSON mode
        try:
            body = await request.json()
        except Exception:
            raise HTTPException(400, "Invalid JSON body")

        raw_content = body.get("content", "")
        cv_filename = body.get("filename", "cv.txt")
        src = body.get("source", source)
        dept = body.get("department", department)
        fmt = body.get("format", "text")

        if fmt == "base64":
            try:
                cv_bytes = base64.b64decode(raw_content)
            except Exception:
                raise HTTPException(400, "Invalid base64 content")
        else:
            cv_bytes = raw_content.encode("utf-8")

    elif file:
        # File upload mode
        cv_bytes = await file.read()
        cv_filename = file.filename or "cv.txt"

    elif content:
        # Form data mode
        cv_bytes = content.encode("utf-8")
        cv_filename = filename or "cv.txt"

    else:
        raise HTTPException(400, "Provide a file, form data (content), or JSON body")

    if not cv_bytes:
        raise HTTPException(400, "Empty CV content")

    # Extract text
    cv_text = await _extract_text(cv_bytes, cv_filename)
    if not cv_text.strip():
        raise HTTPException(422, "Could not extract text from the provided CV")

    # Save the raw file
    file_id = str(uuid.uuid4())
    ext = cv_filename.rsplit(".", 1)[-1].lower() if "." in cv_filename else "txt"
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(file_path, "wb") as f:
        f.write(cv_bytes)

    # Parse and create candidate
    candidate = await _create_candidate_from_cv(cv_text, cv_filename, src, dept, db)

    # Audit log
    db.add(AuditLog(
        action="webhook_cv_received",
        entity_type="candidate",
        entity_id=str(candidate.candidate_id),
        details={
            "filename": cv_filename,
            "source": src,
            "department": dept,
        },
    ))
    await db.commit()

    # Reload with relationships
    result = await db.execute(
        select(Candidate)
        .options(
            selectinload(Candidate.skills),
            selectinload(Candidate.employment_history),
            selectinload(Candidate.sync_status),
            selectinload(Candidate.documents),
        )
        .where(Candidate.candidate_id == candidate.candidate_id)
    )
    candidate = result.scalar_one()

    return JSONResponse(
        status_code=201,
        content={
            "status": "ok",
            "candidate_id": str(candidate.candidate_id),
            "full_name": candidate.full_name,
            "email": candidate.email,
            "phone": candidate.phone,
            "city": candidate.city,
            "career_level": candidate.career_level.value if candidate.career_level else None,
            "department": candidate.department_tag,
            "skills_count": len(candidate.skills),
            "experience_count": len(candidate.employment_history),
            "sync_status": candidate.sync_status.sync_status.value if candidate.sync_status else None,
            "is_duplicate": candidate.is_duplicate_of is not None,
        },
    )


@router.post("/cv/batch")
async def receive_cv_batch(
    request: Request,
    x_api_key: Optional[str] = Header(None),
    db: AsyncSession = Depends(get_db),
):
    """
    Receive multiple CVs in one call.

    POST /api/webhook/cv/batch
    Headers: X-API-Key: your-key
    Body:
    {
      "cvs": [
        {"content": "CV text 1...", "filename": "cv1.pdf", "source": "linkedin"},
        {"content": "CV text 2...", "filename": "cv2.pdf", "source": "referral"}
      ]
    }
    """
    await _verify_api_key(x_api_key, db)

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON body")

    cv_list = body.get("cvs", [])
    if not cv_list:
        raise HTTPException(400, "Provide a 'cvs' array")

    if len(cv_list) > 50:
        raise HTTPException(400, "Maximum 50 CVs per batch")

    results = []
    for i, item in enumerate(cv_list):
        try:
            raw = item.get("content", "")
            fmt = item.get("format", "text")
            fname = item.get("filename", f"cv_{i}.txt")

            if fmt == "base64":
                cv_bytes = base64.b64decode(raw)
            else:
                cv_bytes = raw.encode("utf-8")

            cv_text = await _extract_text(cv_bytes, fname)
            if not cv_text.strip():
                results.append({"index": i, "status": "error", "reason": "Empty text"})
                continue

            candidate = await _create_candidate_from_cv(
                cv_text, fname, item.get("source"), item.get("department"), db
            )
            results.append({
                "index": i,
                "status": "ok",
                "candidate_id": str(candidate.candidate_id),
                "full_name": candidate.full_name,
                "email": candidate.email,
            })
        except Exception as e:
            results.append({"index": i, "status": "error", "reason": str(e)})

    await db.commit()

    ok_count = sum(1 for r in results if r["status"] == "ok")
    return JSONResponse(
        status_code=201 if ok_count > 0 else 422,
        content={
            "total": len(cv_list),
            "processed": ok_count,
            "failed": len(cv_list) - ok_count,
            "results": results,
        },
    )


@router.get("/status")
async def webhook_status(
    x_api_key: Optional[str] = Header(None),
    db: AsyncSession = Depends(get_db),
):
    """Health check for the webhook — requires API key."""
    await _verify_api_key(x_api_key, db)
    return {"status": "ok", "message": "Webhook is active"}


@router.post("/config")
async def configure_webhook_key(
    request: Request,
    x_api_key: Optional[str] = Header(None),
    db: AsyncSession = Depends(get_db),
):
    """
    Set or update the webhook API key.

    POST /api/webhook/config
    Headers: X-API-Key: <admin JWT token> (uses regular auth)
    Body: {"api_key": "your-new-webhook-secret"}
    """
    # For config, we require regular admin auth
    from app.auth import get_current_user, require_permission
    from app.models import User

    # Verify the caller is an admin via the X-API-Key being a JWT token
    from app.auth import decode_token
    if not x_api_key:
        raise HTTPException(401, "Admin auth required")

    try:
        payload = decode_token(x_api_key)
        user_id = payload.get("sub")
    except Exception:
        raise HTTPException(401, "Invalid admin token")

    result = await db.execute(select(User).where(User.user_id == user_id))
    user = result.scalar_one_or_none()
    if not user or not user.is_active:
        raise HTTPException(401, "User not found")

    from app.models import UserRole
    if user.role not in (UserRole.super_admin, UserRole.hr_admin):
        raise HTTPException(403, "Only admins can configure webhook")

    body = await request.json()
    new_key = body.get("api_key", "")
    if not new_key:
        raise HTTPException(400, "Provide an api_key")

    # Save to system settings
    result = await db.execute(
        select(SystemSetting).where(SystemSetting.key == "webhook_api_key")
    )
    setting = result.scalar_one_or_none()
    if setting:
        setting.value = new_key
    else:
        db.add(SystemSetting(key="webhook_api_key", value=new_key))

    await db.commit()
    return {"status": "ok", "message": "Webhook API key updated"}
