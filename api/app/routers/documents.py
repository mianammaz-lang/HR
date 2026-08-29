"""
CV upload and document management endpoints.
Uses built-in parser by default, AI as optional enhancement.
"""
import os
import hashlib
import uuid
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from app.database import get_db
from app.models import Candidate, Document, Skill, EmploymentHistory, SyncStatusRecord, User, AuditLog
from app.schemas import CandidateResponse, DocumentResponse
from app.auth import require_permission
from app.services.cv_parser import BuiltInCVParser
from app.config import UPLOAD_DIR

router = APIRouter(prefix="/api/documents", tags=["Documents"])


@router.post("/upload-cv", response_model=CandidateResponse)
async def upload_cv(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_permission("documents:write")),
):
    """Upload a CV, parse it with built-in parser, and create a candidate record."""
    # Save file
    file_id = str(uuid.uuid4())
    ext = (file.filename or "cv").rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "pdf"
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")

    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # Extract text from CV
    cv_text = ""
    try:
        if ext == "pdf":
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    cv_text += text + "\n"
        elif ext in ("docx", "doc"):
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(content))
            for para in doc.paragraphs:
                if para.text.strip():
                    cv_text += para.text + "\n"
        elif ext == "txt":
            cv_text = content.decode("utf-8", errors="ignore")
        else:
            cv_text = content.decode("utf-8", errors="ignore")
    except Exception as e:
        cv_text = f"CV file: {file.filename}\nError extracting text: {str(e)}"

    if not cv_text.strip():
        cv_text = f"CV file: {file.filename or 'Unknown'}"

    # ── Parse with built-in parser ────────────────────────────────────────
    parser = BuiltInCVParser()
    parsed = parser.parse(cv_text)

    # ── Optional: Try AI enhancement if OpenRouter key is set ──────────────
    from app.config import settings
    if settings.OPENROUTER_API_KEY:
        try:
            from app.services.ai_service import AIService
            ai = AIService()
            ai_parsed = await ai.parse_cv(cv_text)
            # Merge: prefer AI results, fallback to built-in
            for key in ("full_name", "email", "phone", "linkedin_url", "city"):
                if ai_parsed.get(key) and (not parsed.get(key) or parsed[key] == "Unknown"):
                    parsed[key] = ai_parsed[key]
            if ai_parsed.get("skills") and len(ai_parsed["skills"]) > len(parsed.get("skills", [])):
                parsed["skills"] = ai_parsed["skills"]
            if ai_parsed.get("employment_history"):
                parsed["employment_history"] = ai_parsed["employment_history"]
            if ai_parsed.get("career_level"):
                parsed["career_level"] = ai_parsed["career_level"]
        except Exception:
            pass  # Built-in parser already handled it

    # ── Calculate duplicate hash ──────────────────────────────────────────
    email = (parsed.get("email") or "").strip().lower()
    phone = (parsed.get("phone") or "").strip()
    raw_hash = (email + phone) if (email or phone) else ""
    dup_hash = hashlib.sha256(raw_hash.encode()).hexdigest() if raw_hash else None

    # Check for duplicates
    is_duplicate_of = None
    if dup_hash:
        existing = await db.execute(
            select(Candidate).where(Candidate.duplicate_check_hash == dup_hash)
        )
        existing_candidate = existing.scalar_one_or_none()
        if existing_candidate:
            is_duplicate_of = existing_candidate.candidate_id

    # ── Map career level ──────────────────────────────────────────────────
    from app.models import CareerLevel
    career_level_map = {
        "entry": "Entry", "junior": "Entry",
        "mid": "Mid", "intermediate": "Mid",
        "senior": "Senior", "sr": "Senior",
        "lead": "Lead",
        "managerial": "Managerial", "manager": "Managerial", "director": "Managerial",
    }
    cl_raw = (parsed.get("career_level") or "Mid").lower()
    cl = None
    try:
        cl = CareerLevel(career_level_map.get(cl_raw, "Mid"))
    except ValueError:
        cl = CareerLevel.mid

    # ── Create candidate ──────────────────────────────────────────────────
    candidate = Candidate(
        full_name=parsed.get("full_name", "Unknown"),
        email=parsed.get("email"),
        phone=parsed.get("phone"),
        linkedin_url=parsed.get("linkedin_url"),
        city=parsed.get("city"),
        applied_designation=parsed.get("applied_designation"),
        department_tag=parsed.get("department_tag"),
        career_level=cl,
        tags=parsed.get("tags", []),
        duplicate_check_hash=dup_hash,
        is_duplicate_of=is_duplicate_of,
    )
    db.add(candidate)
    await db.flush()

    # Add skills
    for skill_name in parsed.get("skills", []):
        db.add(Skill(candidate_id=candidate.candidate_id, skill_name=skill_name))

    # Add employment history
    for hist in parsed.get("employment_history", []):
        db.add(EmploymentHistory(
            candidate_id=candidate.candidate_id,
            company=hist.get("company", ""),
            title=hist.get("title", ""),
            duration=hist.get("duration", ""),
            description=hist.get("description", ""),
        ))

    # Add document record
    db.add(Document(
        candidate_id=candidate.candidate_id,
        resume_file_url=file_path,
        original_filename=file.filename,
        file_size=len(content),
    ))

    # Add sync status
    db.add(SyncStatusRecord(candidate_id=candidate.candidate_id))

    # Audit log
    db.add(AuditLog(
        user_id=user.user_id,
        action="cv_uploaded",
        entity_type="candidate",
        entity_id=str(candidate.candidate_id),
        details={
            "filename": file.filename,
            "parser": "ai_enhanced" if settings.OPENROUTER_API_KEY else "builtin",
            "extracted_fields": list(parsed.keys()),
            "skills_found": len(parsed.get("skills", [])),
            "experience_found": len(parsed.get("employment_history", [])),
        },
    ))

    await db.flush()

    # Reload with relationships
    result = await db.execute(
        select(Candidate)
        .options(
            selectinload(Candidate.skills),
            selectinload(Candidate.employment_history),
            selectinload(Candidate.scores),
            selectinload(Candidate.sync_status),
            selectinload(Candidate.documents),
        )
        .where(Candidate.candidate_id == candidate.candidate_id)
    )
    candidate = result.scalar_one()
    return CandidateResponse.model_validate(candidate)


@router.get("/{doc_id}/download")
async def download_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_permission("documents:read")),
):
    from fastapi.responses import FileResponse
    result = await db.execute(select(Document).where(Document.doc_id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc or not doc.resume_file_url:
        raise HTTPException(status_code=404, detail="Document not found")
    return FileResponse(
        doc.resume_file_url,
        filename=doc.original_filename or "document",
    )
